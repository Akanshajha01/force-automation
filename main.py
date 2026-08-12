# COMPLETE REVISED `main.p'
# =========================================================
# CALIBRATION REPORT AUTOMATION SYSTEM
# CSIR-NPL FORCE & HARDNESS LABORATORY
# COMPLETE REVISED VERSION
# =========================================================

import threading
import webbrowser
import time
import os
import sys
import socket
import logging
import traceback

from dash import Dash, html, dcc, dash_table
from dash import ctx
import dash
from dash.dependencies import Input, Output, State
import dash_bootstrap_components as dbc
from flask import Flask
from dateutil.relativedelta import relativedelta

import pandas as pd
import numpy as np
import io
import base64

from docx import Document
from docx.shared import Inches
from datetime import datetime

from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows


# =========================================================
# PACKAGING / DEPLOYMENT HELPERS
# =========================================================
# These helpers make the application work correctly both when
# run directly with Python (development) and when packaged into
# a standalone Windows executable with PyInstaller (production).
# They do NOT change any calibration, calculation, or certificate
# logic below.

def resource_path(relative_path):
    """
    Resolve the absolute path to a bundled resource (templates,
    assets, etc.) so it works both:
      1. When running normally with `python main.py`
      2. When running from a PyInstaller-built executable
         (onefile:  files live under sys._MEIPASS)
         (onedir:   files live next to the .exe)
    """
    if getattr(sys, "frozen", False):
        base_path = getattr(sys, "_MEIPASS", os.path.dirname(sys.executable))
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)


def get_output_dir():
    """
    Certificates must be written to a location the user can always
    write to (never into Program Files or PyInstaller's temporary
    extraction folder). We use a folder under the user's Documents.
    """
    docs_dir = os.path.join(
        os.path.expanduser("~"), "Documents", "Force Automation", "Certificates"
    )
    os.makedirs(docs_dir, exist_ok=True)
    return docs_dir


def get_log_dir():
    """Writable, per-user location for application logs."""
    appdata = os.environ.get("APPDATA") or os.path.expanduser("~")
    log_dir = os.path.join(appdata, "Force Automation", "logs")
    os.makedirs(log_dir, exist_ok=True)
    return log_dir


def find_free_port(preferred_port=5080):
    """
    Try the preferred port first (keeps existing 127.0.0.1:5080
    architecture). If it's already in use, fall back to any free
    local port so the app still starts reliably.
    """
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        try:
            s.bind(("127.0.0.1", preferred_port))
            return preferred_port
        except OSError:
            pass
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return s.getsockname()[1]


LOG_DIR = get_log_dir()
LOG_FILE = os.path.join(LOG_DIR, "force_automation.log")

logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)


def _log_uncaught_exceptions(exc_type, exc_value, exc_traceback):
    """
    In a windowed (--noconsole) build there is no terminal to show
    errors, so uncaught exceptions must be written to the log file
    instead of disappearing silently.
    """
    if issubclass(exc_type, KeyboardInterrupt):
        sys.__excepthook__(exc_type, exc_value, exc_traceback)
        return
    logging.critical(
        "Uncaught exception:\n%s",
        "".join(traceback.format_exception(exc_type, exc_value, exc_traceback)),
    )


sys.excepthook = _log_uncaught_exceptions

OUTPUT_DIR = get_output_dir()


# =========================================================
# SERVER
# =========================================================

server = Flask(__name__)

app = Dash(
    __name__,
    server=server,
    suppress_callback_exceptions=True,
    external_stylesheets=[dbc.themes.BOOTSTRAP],
    assets_folder=resource_path("assets"),
)

app.title = "Calibration Report Automation"



# ==========================================
# POSITION CONFIGURATION
# ==========================================

POSITION_MAP = {

    "loadcell": {

        "s1": "Position 0° series 1",
        "s2": "Position 0° series 2",
        "s3": "Position 120° series 3",
        "s4": "Position 240° series 4"

    },

    "provingring": {

        "s1": "Position 0° series 1",
        "s2": "Position 0° series 2",
        "s3": "Position 180° series 3",
        "s4": "Position 360° series 4"

    }
}


# ==========================================
# READING-SETTING DEFAULTS
# (Proving Ring -> Small Pointer Setting, Load Cell -> No Load Output)
# Kept as named constants so the default text is configurable in one place.
# ==========================================

DEFAULT_SMALL_POINTER_SETTING = "Small pointer at 1 (5 revolutions)"
DEFAULT_NO_LOAD_OUTPUT = "0"

# =========================================================
# RESOLUTION UNIT (NEW)
# Single source of truth for the Resolution-unit dropdown
# (shared by Load Cell + Proving Ring forms) and its default.
# The Load Cell "Digital Indicator Reading in ..." text in the
# Results section is driven directly from this value -- it is
# never hard-coded again anywhere else.
# =========================================================
RESOLUTION_UNIT_OPTIONS = [
    {"label": "div", "value": "div"},
    {"label": "kN", "value": "kN"},
    {"label": "count", "value": "count"},
    {"label": "mV DC", "value": "mV DC"}
]

DEFAULT_RESOLUTION_UNIT = "div"


FORCE_MACHINE_MAP = {

    "1": "1 kN Dead Weight Force Machine",
    "2": "5 kN Dead Weight Force Machine",
    "3": "50 kN Dead Weight Force Machine",
    "4": "100 kN Dead Weight Force Machine",
    "5": "1 MN FSM Force Machine",
    "6": "1000 kN Lever Multiplication Force Machine",
    "7": "1000 kN HMS Force Machine",
    "8": "200 kN Force Machine",
    "9": "3000 kN Force Machine"

}

BMC_VALUES = {

    "1": 0.008,
    "2": 0.008,
    "3": 0.007,
    "4": 0.002,
    "5": 0.003,
    "6": 0.009,
    "7": 0.020,
    "8": 0.025,
    "9": 0.050

}
# =========================================================
# DEFAULT TABLE
# =========================================================

default_df = pd.DataFrame({

    "Applied Force kN": [0],
    "Position 0° series 1": [0],
    "Position 0° series 2": [0],
    "Position 120° series 3": [0],
    "Position 240° series 4": [0],
    "Average 1,3,4": [0]

})


def format_force(value):
    value = float(value)

    # If integer, show without decimals
    if value.is_integer():
        return f"{int(value)}"

    # Otherwise show up to 2 decimals
    return f"{value:.2f}"
# =========================================================
# STYLES
# =========================================================

CARD_STYLE = {
    "borderRadius": "14px",
    "boxShadow": "0px 3px 12px rgba(0,0,0,0.12)",
    "marginBottom": "14px",
    "border": "none"
}

CARD_HEADER_STYLE = {
    "background": "#003F73",
    "color": "white",
    "fontWeight": "600",
    "fontSize": "18px",
    "padding": "10px 16px"
}

INPUT_STYLE = {
    "height": "38px",
    "fontSize": "14px",
    "borderRadius": "8px"
}

BUTTON_STYLE = {
    "borderRadius": "10px",
    "fontWeight": "600",
    "padding": "10px 18px",
    "fontSize": "15px"
}

LABEL_STYLE = {
    "fontWeight": "600",
    "fontSize": "14px"
}

# =========================================================
# READING-SETTING FIELD BUILDER
# (Small Pointer Setting for Proving Ring / No Load Output for Load Cell)
# Shared by the initial page layout AND the dynamic callback so the
# field is never dependent on a callback round-trip to become visible.
# =========================================================

def build_reading_setting_field(cert_type):

    if cert_type == "provingring":

        return html.Div([
            html.Label("Small Pointer Setting", style=LABEL_STYLE),
            dbc.Input(
                id="small_pointer_setting",
                type="text",
                value=DEFAULT_SMALL_POINTER_SETTING,
                style={
                    "width": "100%",
                    "fontSize": "14px",
                    "height": "42px"
                }
            )
        ])

    # Default / loadcell case
    return html.Div([
        html.Label("No Load Output", style=LABEL_STYLE),
        dbc.InputGroup([

            dcc.Input(
                id="no_load_output",
                type="number",
                className="form-control",
                style=INPUT_STYLE,
                value=0
            ),

            # Read-only unit badge -- NEVER typed by the user. It is kept
            # in sync with the Resolution Unit dropdown by the
            # sync_no_load_output_unit callback below (single source of
            # truth = the "resolution_unit" component's value).
            dbc.InputGroupText(
                id="no_load_output_unit_display",
                children=DEFAULT_RESOLUTION_UNIT
            )

        ])
    ])


@app.callback(
    Output("no_load_output_unit_display", "children"),
    Input("resolution_unit", "value"),
    prevent_initial_call=False
)
def sync_no_load_output_unit(resolution_unit):
    """
    Keeps the (read-only) No Load Output unit badge identical to the
    Resolution Unit dropdown at all times. This is the single place
    that drives that badge -- the user never types this unit manually.
    Only fires when the Load Cell form (and therefore the badge) is on
    screen; suppress_callback_exceptions=True on the app already allows
    this safely for the Proving Ring form, where the badge is absent.
    """
    return resolution_unit or DEFAULT_RESOLUTION_UNIT

# =========================================================
# ISO FUNCTIONS
# =========================================================

def compute_repeatability(values):

    values = np.array(values, dtype=float)

    mean_val = np.mean(values)

    std_val = np.std(values, ddof=1)

    if mean_val == 0:
        return 0

    return abs(std_val / mean_val) * 100

# =========================================================

def compute_reproducibility(series1, series2, series3, series4):

    arr = np.vstack([
        series1,
        series2,
        series3,
        series4
    ])

    max_vals = np.max(arr, axis=0)

    min_vals = np.min(arr, axis=0)

    avg_vals = np.mean(arr, axis=0)

    repro = np.max(
        np.abs(max_vals - min_vals) / avg_vals
    ) * 100

    return repro

# =========================================================

def compute_reversibility(ascending, descending):

    ascending = np.array(ascending, dtype=float)

    descending = np.array(descending, dtype=float)

    diff = np.abs(ascending - descending)

    max_output = np.max(ascending)

    if max_output == 0:
        return 0

    return np.max(diff) / max_output * 100

# =========================================================

def compute_interpolation_error(applied_force, indicated_force):

    applied_force = np.array(applied_force, dtype=float)

    indicated_force = np.array(indicated_force, dtype=float)

    if len(applied_force) < 4:

        return 0, np.array([0,0,0,0])

    coeffs = np.polyfit(
        indicated_force,
        applied_force,
        3
    )

    predicted = np.polyval(coeffs, indicated_force)

    error = np.abs(predicted - applied_force)

    max_force = np.max(applied_force)

    interpolation_error = (
        np.max(error) / max_force
    ) * 100

    return interpolation_error, coeffs


# =========================================================

def compute_uncertainty(
    repeatability,
    reproducibility,
    reversibility,
    interpolation,
    standard_uncertainty
):

    combined = np.sqrt(
        repeatability**2 +
        reproducibility**2 +
        reversibility**2 +
        interpolation**2 +
        standard_uncertainty**2
    )

    expanded = combined * 2

    return expanded

# =========================================================
def generate_equation(coeffs):

    if coeffs is None:
        return ""

    def fmt(value, variable=""):
        sign = "+" if value >= 0 else "-"
        return f" {sign} {abs(value):.8E}{variable}"

    return (
        f"{coeffs[0]:.8E}X³"
        + fmt(coeffs[1], "X²")
        + fmt(coeffs[2], "X")
        + fmt(coeffs[3])
    )

# =========================================================

def full_iso376_analysis(df, standard_uncertainty=0.002):



    if "Position 180° series 3" in df.columns:

        s1 = df["Position 0° series 1"].astype(float)
        s2 = df["Position 0° series 2"].astype(float)
        s3 = df["Position 180° series 3"].astype(float)
        s4 = df["Position 360° series 4"].astype(float)

    else:

        s1 = df["Position 0° series 1"].astype(float)
        s2 = df["Position 0° series 2"].astype(float)
        s3 = df["Position 120° series 3"].astype(float)
        s4 = df["Position 240° series 4"].astype(float)

    average = (s1 + s3 + s4) / 3

    df.iloc[:,5] = np.round(average, 6)

    repeatability = compute_repeatability(s1)

    reproducibility = compute_reproducibility(
        s1,
        s2,
        s3,
        s4
    )

    reversibility = compute_reversibility(s1, s2)
    force = df["Applied Force kN"].astype(float)

    interpolation_error, coeffs = compute_interpolation_error(
        force,
        average
    )
    
    equation = generate_equation(coeffs)

    uncertainty = compute_uncertainty(
        repeatability,
        reproducibility,
        reversibility,
        interpolation_error,
        standard_uncertainty
    )

    worst = max([
        repeatability,
        reproducibility,
        reversibility,
        interpolation_error,
        uncertainty
    ])

    final_class = determine_class(worst)

    equation = generate_equation(coeffs)

    return {
        "updated_dataframe": df,
        "repeatability": repeatability,
        "reproducibility": reproducibility,
        "reversibility": reversibility,
        "interpolation_error": interpolation_error,
        "uncertainty": uncertainty,
        "classification": final_class,
        "coefficients": coeffs,
        "equation": equation
    }



# =========================================================
# BACKEND PROCESSING FUNCTION
# =========================================================

def determine_iso_class(
        repr_val,
        rep,
        reso,
        zero,
        creep,
        interp,
        BMC,
        final_w):

    classes = {

        "00": {
            "repr":0.05,
            "rep":0.025,
            "zero":0.012,
            "creep":0.025,
            "interp":0.025,
            "BMC": 0.01
        },

        "0.5":{
            "repr":0.10,
            "rep":0.05,
            "zero":0.025,
            "creep":0.05,
            "interp":0.05,
            "BMC": 0.02
        },

        "1":{
            "repr":0.20,
            "rep":0.10,
            "zero":0.05,
            "creep":0.10,
            "interp":0.1,
            "BMC": 0.05
        },

        "2":{
            "repr":0.40,
            "rep":0.20,
            "zero":0.10,
            "creep":0.20,
            "interp":0.2,
            "BMC": 0.1
        }
    }

    for cls in ["00","0.5","1","2"]:

        lim = classes[cls]

        if (
            repr_val <= lim["repr"]
            and rep <= lim["rep"]
            and zero <= lim["zero"]
            and creep <= lim["creep"]
            and interp <= lim["interp"]
            and BMC <= lim["BMC"]
          
        ):
            return cls

    return "Out of Class"


# =====================================================
# DISPLAY-ONLY FORMATTING HELPER
# =====================================================
# NOTE: This function is used PURELY to decide how many
# decimal places are shown in the generated certificate.
# It does not participate in, and must never be used for,
# any calculation, interpolation, classification or
# uncertainty logic.
#
# Logic:
#   Resolution = 1        -> 0 decimal places
#   Resolution = 0.1      -> 1 decimal place
#   Resolution = 0.01     -> 2 decimal places
#   Resolution = 0.001    -> 3 decimal places
#   Resolution = 0.0001   -> 4 decimal places
#   Resolution = 0.00001  -> 5 decimal places
#
# The precision is detected dynamically from the text of
# the Resolution field itself (number of digits after the
# decimal point), so it is never hardcoded.
def get_decimal_places_from_resolution(resolution_value):

    if resolution_value is None:
        return None

    text = str(resolution_value).strip()

    if text in ("", "nan", "None"):
        return None

    # Reject anything that isn't a valid number so we fall
    # back to the previous (pre-existing) behaviour safely.
    try:
        float(text)
    except (TypeError, ValueError):
        return None

    if "." in text:
        decimals = len(text.split(".")[1].rstrip())
    else:
        decimals = 0

    return decimals


def process_calibration_data(input_data):

    df = pd.DataFrame(input_data["table_data"]).copy()

    print("\n===== DATAFRAME COLUMNS =====")
    print(df.columns.tolist())
    print("=============================\n")

    print(df.head())

    # ==========================================
    # CLEAN DATA BEFORE CONVERSION
    # ==========================================

    for col in df.columns:

        df[col] = (
            df[col]
            .astype(str)
            .str.replace("\r", "", regex=False)
            .str.strip()
        )

        df[col] = pd.to_numeric(
            df[col],
            errors="coerce"
        )

    division_value = float(
    input_data.get("div_value") or 0.1
         )


    resolution_div = float(
            input_data.get("resolution") or 0
        )
    
    temperature = float(input_data["temperature"] or 23)

    temp_factor = (
        1 - 0.00027 * (temperature - 23)
    )

    cert_type = input_data.get("certificate_type", "loadcell")
    ring_type = input_data.get("ring_type", "")
    
    if cert_type == "provingring":

        expected_cols = [

            "Applied Force kN",
            "Position 0° series 1",
            "Position 0° series 2",
            "Position 180° series 3",
            "Position 360° series 4",
            "Average 1,3,4"

        ]

    else:

        expected_cols = [

            "Applied Force kN",
            "Position 0° series 1",
            "Position 0° series 2",
            "Position 120° series 3",
            "Position 240° series 4",
            "Average 1,3,4"

        ]

    df = df[[c for c in expected_cols if c in df.columns]]

    s1_col = POSITION_MAP[cert_type]["s1"]
    s2_col = POSITION_MAP[cert_type]["s2"]
    s3_col = POSITION_MAP[cert_type]["s3"]
    s4_col = POSITION_MAP[cert_type]["s4"]
        
    if s3_col not in df.columns:

        if "Position 180° series 3" in df.columns:
            s3_col = "Position 180° series 3"

    if s4_col not in df.columns:

        if "Position 360° series 4" in df.columns:
            s4_col = "Position 360° series 4"
    # =====================================================
    # TEMPERATURE COMPENSATION
    # Apply only for Proving Ring / Bow Dynamometer
    # =====================================================

    if cert_type == "provingring":

        df["S1_TC"] = df[s1_col] * temp_factor
        df["S2_TC"] = df[s2_col] * temp_factor
        df["S3_TC"] = df[s3_col] * temp_factor
        df["S4_TC"] = df[s4_col] * temp_factor

    else:
        # Load Cell -> No temperature compensation

        df["S1_TC"] = df[s1_col]
        df["S2_TC"] = df[s2_col]
        df["S3_TC"] = df[s3_col]
        df["S4_TC"] = df[s4_col]

    # =====================================================
    # FULL SCALE ROW
    # =====================================================
    force_col = "Applied Force kN"
    valid_force_rows = df[
        pd.to_numeric(df[force_col], errors="coerce") > 0
    ]

    if valid_force_rows.empty:

        raise ValueError(
            "No positive force values found."
        )

    full_row = valid_force_rows.loc[
        valid_force_rows[force_col].idxmax()
    ]
    # =====================================================
    # FINAL ZERO ROW
    # =====================================================

    zero_candidates = df[
        pd.to_numeric(df[force_col], errors="coerce") == 0
    ]

    if len(zero_candidates) == 0:

        raise ValueError(
            "Calibration table must contain at least one zero-force row."
        )

    elif len(zero_candidates) > 1:

        zero_row = zero_candidates.iloc[-1]

    else:

        zero_row = zero_candidates.iloc[0]

    # =====================================================
    # ZERO ERROR
    # =====================================================

    zero_ratios = []

    for col in ["S1_TC","S2_TC","S3_TC","S4_TC"]:

        fs = float(full_row[col])

        zr = float(zero_row[col])

        if fs > 0:

            zero_ratios.append(
                abs(zr)/fs*100
            )

    zero_error = max(zero_ratios)

    # =====================================================
    # CREEP
    # =====================================================

    try:

        final_zero_s4 = float(zero_row["S4_TC"])

        max_load_s4 = float(full_row["S4_TC"])

        creep = (
            abs(final_zero_s4 - 0.0)
            /
            max_load_s4
        ) * 100

    except:

        creep = 0


    calc_df = df[
        pd.to_numeric(df[force_col], errors="coerce") > 0
    ].copy()

    # ==========================================
    # DETECT DECIMALS FROM SERIES 1 INPUT
    # ==========================================
    # (UNCHANGED - this still drives REP/REPR/error/BMC
    # display precision exactly as before, per the
    # "do not touch calculation-adjacent logic" instruction.)

# =====================================================
# DETECT DECIMALS FROM ALL SERIES-1 VALUES
# =====================================================

    reading_decimals = 2

    try:

        original_df = pd.DataFrame(input_data["table_data"])

        decimal_counts = []

        for val in original_df[s1_col]:

            txt = str(val).strip()

            if txt in ["", "nan", "None"]:
                continue

            if "." in txt:

                decimal_counts.append(
                    len(txt.split(".")[1])
                )

            else:

                decimal_counts.append(0)

        if decimal_counts:

            reading_decimals = max(decimal_counts)

    except:

        reading_decimals = 2

    print("READING_DECIMALS =", reading_decimals)
    print("Detected reading decimals =", reading_decimals)

    # ==========================================
    # RESOLUTION-DRIVEN DISPLAY PRECISION (NEW)
    # ==========================================
    # This is the ONLY variable that should govern the
    # decimal formatting of instrument readings shown in
    # the certificate (Position Series readings, Average
    # values, No Load Output, Calibration Signal). It is
    # detected dynamically from the Resolution field and
    # is display-only - it never feeds any calculation,
    # interpolation, classification or uncertainty logic.
    #
    # If Resolution is missing/invalid, we fall back to
    # the previously-existing reading_decimals detection
    # above so behaviour is unchanged for old data.

    resolution_display_decimals = get_decimal_places_from_resolution(
        input_data.get("resolution")
    )

    if resolution_display_decimals is None:
        resolution_display_decimals = reading_decimals

    print(
        "RESOLUTION_DISPLAY_DECIMALS =",
        resolution_display_decimals
    )

    avg34_list = []
    avg134_list = []

    rep_list = []
    repr_list = []
    reso_list = []

    interp_list = []

    w_list = []
    final_w_list = []

    class_list = []

    BMC = float(
        input_data.get("bmc_value", 0.002)
    )
    # =====================================================
    # CALCULATIONS
    # =====================================================
# Full scale force

    full_scale = float(
        calc_df[force_col].max()
    )
# =====================================================
# LOAD CELL INTERPOLATION SERIES SELECTION
# =====================================================

    interp_series_selected = np.zeros(len(calc_df))

    if cert_type == "loadcell":

        print("========== DEBUG ==========")
        print("cert_type =", repr(cert_type))
        print("===========================")

        force_fit_data = calc_df[force_col].astype(float).values

        avg_fit_data = (
            calc_df["S1_TC"]
            + calc_df["S3_TC"]
            + calc_df["S4_TC"]
        ) / 3

        avg_fit_data = avg_fit_data.values

        # ----------------------------------
        # FIT 1 : Force -> Average
        # ----------------------------------

        coeff_f_to_a = np.polyfit(
            force_fit_data,
            avg_fit_data,
            3
        )

        fitted_avg = np.polyval(
            coeff_f_to_a,
            force_fit_data
        )

        err_fit1_percent = np.where(

            avg_fit_data != 0,

            (
                (avg_fit_data - fitted_avg)
                /
                avg_fit_data
            ) * 100,

            0

        )

        fit1_max = np.max(
            np.abs(err_fit1_percent)
        )

        # ----------------------------------
        # FIT 2 : Average -> Force
        # ----------------------------------

        coeff_a_to_f = np.polyfit(
            avg_fit_data,
            force_fit_data,
            3
        )

        fitted_force = np.polyval(
            coeff_a_to_f,
            avg_fit_data
        )

        err_fit2_percent = np.where(

            force_fit_data != 0,

            (
                (force_fit_data - fitted_force)
                /
                force_fit_data
            ) * 100,

            0

        )

        fit2_max = np.max(
            np.abs(err_fit2_percent)
        )

        print("\n========== INTERPOLATION ==========")

        print(
            "Fit1 Max Error (%) =",
            fit1_max
        )

        print(
            "Fit2 Max Error (%) =",
            fit2_max
        )

        # ----------------------------------
        # SELECT WORSE FIT
        # ----------------------------------

        if fit1_max >= fit2_max:

            interp_series_selected = np.abs(
                err_fit1_percent
            )

            print(
                "Using Force -> Average fit series"
            )

        else:

            interp_series_selected = np.abs(
                err_fit2_percent
            )

            print(
                "Using Average -> Force fit series"
            )

        print("===================================\n")
    # Full scale indication

    avg134_fs = (
        float(full_row["S1_TC"])
        + float(full_row["S3_TC"])
        + float(full_row["S4_TC"])
    ) / 3

    # Force equivalent of one division

    if avg134_fs > 0:

        resolution_force = (
            resolution_div
            * full_scale
            / avg134_fs
        )

    else:

        resolution_force = 0


    force_to_avg_eq = ""
    avg_to_force_eq = ""

    for idx, (_, row) in enumerate(
        calc_df.iterrows()
    ):
# ==========================================
# LOAD CELL INTERPOLATION ERROR
# ==========================================

        if cert_type == "loadcell":

                force_fit_data = calc_df[force_col].astype(float).values

                avg_fit_data = (
                    calc_df["S1_TC"]
                    + calc_df["S3_TC"]
                    + calc_df["S4_TC"]
                ) / 3

                avg_fit_data = avg_fit_data.values

                # Force -> Average

                coeff_f_to_a = np.polyfit(
                    force_fit_data,
                    avg_fit_data,
                    3
                )

                # Average -> Force

                coeff_a_to_f = np.polyfit(
                    avg_fit_data,
                    force_fit_data,
                    3
                )

        def format_term(value, variable=""):
            sign = "+" if value >= 0 else "-"
            return f" {sign} {abs(value):.4E}{variable}"

        if cert_type == "loadcell":

          force_to_avg_eq = (
            f"{coeff_f_to_a[0]:.4E}F³"
            + format_term(coeff_f_to_a[1], "F²")
            + format_term(coeff_f_to_a[2], "F")
            + format_term(coeff_f_to_a[3])
            )

          avg_to_force_eq = (
            f"{coeff_a_to_f[0]:.4E}X³"
            + format_term(coeff_a_to_f[1], "X²")
            + format_term(coeff_a_to_f[2], "X")
            + format_term(coeff_a_to_f[3])
            )

        else:

            force_to_avg_eq = ""
            avg_to_force_eq = ""
        
        if cert_type == "loadcell":

            print("\n========== INTERPOLATION RESULTS ==========")

            print("\nForward Interpolation (Force -> Average)")
            print(force_to_avg_eq)
            print(f"Maximum Error = {fit1_max:.6f} %")

            print("\nBackward Interpolation (Average -> Force)")
            print(avg_to_force_eq)
            print(f"Maximum Error = {fit2_max:.6f} %")

            if fit1_max >= fit2_max:
                print("\nSelected: Forward Interpolation")
            else:
                print("\nSelected: Backward Interpolation")

            print("==========================================")


        if cert_type == "loadcell":

            print("\n=== FORCE -> OUTPUT FIT ===")
            print(force_to_avg_eq)

            print("\n=== OUTPUT -> FORCE FIT ===")

            print(avg_to_force_eq)
            
        s1 = float(row["S1_TC"])
        s2 = float(row["S2_TC"])
        s3 = float(row["S3_TC"])
        s4 = float(row["S4_TC"])

        # ------------------------------------------calibraty
        # Average values
        # ------------------------------------------

        avg34_name = f"AVG({s3_col},{s4_col})"

        avg134_name = f"AVG(0°,{s3_col.split()[1]},{s4_col.split()[1]})"


        avg34 = (s3 + s4) / 2

        avg134 = (s1 + s3 + s4) / 3



        # ------------------------------------------
        # Interpolation Error
        # ------------------------------------------

        if cert_type == "loadcell":

            interp_error = float(
                interp_series_selected[idx]
            )

        else:

            interp_error = 0
        # ------------------------------------------
        # Repeatability
        # ------------------------------------------

        if (s1 + s2) == 0:

            rep = 0

        else:

            rep = (
                2 *
                (
                    max(s1, s2)
                    - min(s1, s2)
                )
                /
                (s1 + s2)
            ) * 100
            # Round to 2 decimals BEFORE classification
        rep = round(rep, 2)
        # ------------------------------------------
        # Reproducibility
        # ------------------------------------------

        if avg134 == 0:

            repr_val = 0

        else:

            repr_val = (
                (
                    max(s1, s3, s4)
                    -
                    min(s1, s3, s4)
                )
                /
                avg134
            ) * 100
            # Round to 2 decimals BEFORE classification
        repr_val = round(repr_val, 2)
        # ------------------------------------------
        # Resolution
        # ------------------------------------------

# ------------------------------------------
# Resolution
# ------------------------------------------


        if avg134 == 0:

            reso = 0

        else:

            reso = (
                resolution_div
                /
                avg134
            ) * 100


        # ------------------------------------------
        # Uncertainty
        # ------------------------------------------

        W = np.sqrt(

            (repr_val**2) / 2 +

            (rep**2) / 3 +

            (reso**2) / 3 +

            (zero_error**2) / 3 +

            (0**2) / 3 +

            (creep**2) / 3 +

            (interp_error**2) / 6

        )

        final_w = np.sqrt(

            W**2 +

            BMC**2

        )

        iso_class = determine_iso_class(

            repr_val,

            rep,

            reso,

            zero_error,

            creep,

            interp_error,

            BMC,

            final_w

        )

        avg34_list.append(avg34)
        avg134_list.append(avg134)

        rep_list.append(rep)
        repr_list.append(repr_val)

        reso_list.append(reso)

        interp_list.append(interp_error)
        w_list.append(W)

        final_w_list.append(final_w)

        class_list.append(iso_class)
        print(
            f"Force={row[force_col]:.0f} "
            f"REP={rep:.4f} "
            f"REPR={repr_val:.4f} "
            f"RESO={reso:.4f} "
            f"ZERO={zero_error:.4f} "
            f"CREEP={creep:.4f} "
            f"FINAL_W={final_w:.4f} "
            f"CLASS={iso_class}"
        )
            # =====================================================
    # ASSIGN COLUMNS AFTER LOOP
    # =====================================================

    calc_df[avg34_name] = avg34_list
    calc_df[avg134_name] = avg134_list

    calc_df["REP_%"] = rep_list

    calc_df["REPR_%"] = repr_list

    calc_df["RESO_%"] = reso_list


    calc_df["ZERO_%"] = zero_error

    calc_df["CREEP_%"] = creep

    calc_df["INTERP_%"] = interp_list

    calc_df["W_INITIAL_%"] = w_list

    calc_df["BMC_%"] = BMC

    calc_df["FINAL_W_%"] = final_w_list

    calc_df["ISO_CLASS"] = class_list

    # =====================================================
    # DISPLAY FORMATTING
    # =====================================================

    if cert_type == "loadcell":

        display_s1 = s1_col
        display_s2 = s2_col
        display_s3 = s3_col
        display_s4 = s4_col

    else:

        display_s1 = s1_col + " TC"
        display_s2 = s2_col + " TC"
        display_s3 = s3_col + " TC"
        display_s4 = s4_col + " TC"

    if cert_type == "provingring":

        calc_df.rename(columns={

            "S1_TC": display_s1,
            "S2_TC": display_s2,
            "S3_TC": display_s3,
            "S4_TC": display_s4

        }, inplace=True)
    # =====================================================
    # FORMAT INDICATION COLUMNS
    # =====================================================

    if cert_type == "loadcell":

        display_decimals = resolution_display_decimals

        calc_df[display_s1] = calc_df[display_s1].round(display_decimals)
        calc_df[display_s2] = calc_df[display_s2].round(display_decimals)
        calc_df[display_s3] = calc_df[display_s3].round(display_decimals)
        calc_df[display_s4] = calc_df[display_s4].round(display_decimals)

    else:

        calc_df[display_s1] = calc_df[display_s1].round(resolution_display_decimals)
        calc_df[display_s2] = calc_df[display_s2].round(resolution_display_decimals)
        calc_df[display_s3] = calc_df[display_s3].round(resolution_display_decimals)
        calc_df[display_s4] = calc_df[display_s4].round(resolution_display_decimals)


    if cert_type == "loadcell":

        display_decimals = resolution_display_decimals

        calc_df[avg34_name] = calc_df[avg34_name].round(
            display_decimals
        )

        calc_df[avg134_name] = calc_df[avg134_name].round(
            display_decimals
        )

    else:

        calc_df[avg34_name] = calc_df[avg34_name].round(resolution_display_decimals)
        calc_df[avg134_name] = calc_df[avg134_name].round(resolution_display_decimals)

    # ==========================================
    # DYNAMIC ERROR DISPLAY PRECISION
    # ==========================================

    if cert_type == "loadcell" and reading_decimals > 2:

        rep_decimals = reading_decimals

        error_decimals = reading_decimals + 2

        bmc_decimals = reading_decimals + 1

        final_w_decimals = reading_decimals

    else:

        # KEEP EXISTING BEHAVIOUR

        rep_decimals = 2

        error_decimals = 4

        bmc_decimals = 3

        final_w_decimals = 2


    calc_df["REP_%"] = calc_df["REP_%"].round(rep_decimals)

    calc_df["REPR_%"] = calc_df["REPR_%"].round(rep_decimals)

    calc_df["RESO_%"] = calc_df["RESO_%"].round(error_decimals)

    calc_df["ZERO_%"] = calc_df["ZERO_%"].round(error_decimals)

    calc_df["CREEP_%"] = calc_df["CREEP_%"].round(error_decimals)

    calc_df["INTERP_%"] = calc_df["INTERP_%"].round(error_decimals)

    calc_df["W_INITIAL_%"] = calc_df["W_INITIAL_%"].round(error_decimals)

    calc_df["BMC_%"] = calc_df["BMC_%"].round(bmc_decimals)

    calc_df["FINAL_W_%"] = calc_df["FINAL_W_%"].round(final_w_decimals)

    if cert_type == "loadcell":

        calc_df[display_s1] = calc_df[display_s1].map(lambda x: f"{x:.{resolution_display_decimals}f}")
        calc_df[display_s2] = calc_df[display_s2].map(lambda x: f"{x:.{resolution_display_decimals}f}")
        calc_df[display_s3] = calc_df[display_s3].map(lambda x: f"{x:.{resolution_display_decimals}f}")
        calc_df[display_s4] = calc_df[display_s4].map(lambda x: f"{x:.{resolution_display_decimals}f}")
        calc_df[avg134_name] = calc_df[avg134_name].map(lambda x: f"{x:.{resolution_display_decimals}f}")

    # ----------------------------------------------------
    # Export table (keep BOTH zero rows)
    # ----------------------------------------------------

# --------------------------------------------------------
# EXPORT TABLE (retain both zero-force rows)
# --------------------------------------------------------

    export_df = df.copy()

    # Calculate average
    export_df["Average 1,3,4"] = (
        export_df["S1_TC"] +
        export_df["S3_TC"] +
        export_df["S4_TC"]
    ) / 3

    # Keep only required columns
    export_df = export_df[[
        "Applied Force kN",
        "S1_TC",
        "S2_TC",
        "S3_TC",
        "S4_TC",
        "Average 1,3,4"
    ]]

    # Rename columns
    export_df.columns = [
        "Applied Force kN",
        "Position 0° series 1",
        "Position 0° series 2",
        POSITION_MAP[cert_type]["s3"],
        POSITION_MAP[cert_type]["s4"],
        "Average 1,3,4"
    ]

    # --------------------------------------------------------
    # Format readings
    # --------------------------------------------------------

    dec = resolution_display_decimals

    for col in export_df.columns[1:5]:

        export_df[col] = (
            export_df[col]
            .astype(float)
            .map(lambda x: f"{x:.{dec}f}")
        )

    # Format average as STRING
    export_df["Average 1,3,4"] = (
        export_df["Average 1,3,4"]
        .astype(float)
        .map(lambda x: f"{x:.{dec}f}")
    )

    # Replace average with "-" for zero-force rows
    zero_mask = export_df["Applied Force kN"] == 0

    export_df.loc[
        zero_mask,
        "Average 1,3,4"
    ] = "-"
    
    print("EXPORT_DF COLUMNS BEFORE RENAME")
    print(export_df.columns.tolist())
    print("COUNT =", len(export_df.columns))


    print(export_df.columns.tolist())
    # =====================================================
    # ISO CLASSIFICATION RANGE GENERATION
    # =====================================================

    unc_map = {
        "00": "±0.03%",
        "0.5": "±0.06%",
        "1": "±0.12%",
        "2": "±0.24%"
    }

    forces = calc_df.iloc[:, 0].astype(float).tolist()

    filtered_data = []

    for force, cls in zip(forces, class_list):

        if force <= 0:
            continue

        filtered_data.append((force, cls))


    iso_limits = {

        "00": {
            "repr":0.05,
            "rep":0.025,
            "zero":0.012,
            "creep":0.025,
            "interp":0.025,
            "BMC": 0.01
            
        },

        "0.5":{
            "repr":0.10,
            "rep":0.05,
            "zero":0.025,
            "creep":0.05,
            "interp":0.05,
            "BMC": 0.02
        },

        "1":{
            "repr":0.20,
            "rep":0.10,
            "zero":0.05,
            "creep":0.10,
            "interp":0.1,
            "BMC": 0.05
        },

        "2":{
            "repr":0.40,
            "rep":0.20,
            "zero":0.10,
            "creep":0.20,
            "interp":0.2,
            "BMC": 0.1
        }
    }

    unc_map = {
        "00":"±0.03%",
        "0.5":"±0.06%",
        "1":"±0.12%",
        "2":"±0.24%"
    }

    class_ranges = []

    force_col = calc_df.columns[0]

    resolution_div = float(
        input_data.get("resolution") or 0
    )

    full_scale = float(
        calc_df[force_col].max()
    )

    multiplier_map = {
        "00": 4000,
        "0.5": 2000,
        "1": 1000,
        "2": 500
    }

    for target_class in ["00", "0.5", "1", "2"]:

        lim = iso_limits[target_class]

        sorted_df = calc_df.sort_values(
            force_col,
            ascending=False
        )

        lowest_force = None
        rows_in_range = []          # <-- collect every row that qualifies for this class

        for _, row in sorted_df.iterrows():

            force_value = float(row[force_col])

            minimum_force_required = max(
                multiplier_map[target_class] * resolution_force,
                0.02 * full_scale
            )

            if force_value < minimum_force_required:
                break

            ok = (
                row["REPR_%"] <= lim["repr"] and
                row["REP_%"] <= lim["rep"] and
                row["ZERO_%"] <= lim["zero"] and
                row["CREEP_%"] <= lim["creep"] and
                row["INTERP_%"] <= lim["interp"] and
                row["BMC_%"] <= lim["BMC"]
            )

            if ok:
                lowest_force = force_value
                rows_in_range.append(row)     # <-- keep the whole row, not just the force
            else:
                break

        # Ignore meaningless single-point ranges
        if lowest_force is None or len(rows_in_range) < 2:
            continue

        # worst-case (max) FINAL_W_% across every point actually in this class's range
        max_uncertainty = max(r["FINAL_W_%"] for r in rows_in_range)

        iso_unc_limits = {
            "00": 0.03,
            "0.5": 0.06,
            "1": 0.12,
            "2": 0.24
        }

        display_uncertainty = max(
            max_uncertainty,
            iso_unc_limits[target_class]
        )

# ISO requirement:
# Class 00 / 0.5 must extend to at least 50% FS

        if target_class in ["00", "0.5"]:

            if lowest_force > 0.5 * full_scale:

                print(
                    f"Skipping Class {target_class} "
                    f"because range only extends to "
                    f"{lowest_force} kN "
                    f"(must reach <= {0.5*full_scale} kN)"
                )

                continue

        class_ranges.append({
            "class": target_class,
            "from_force": float(full_scale),
            "to_force": float(lowest_force),
            "uncertainty": f"±{display_uncertainty:.2f}%"
        })


    filtered = []

    for r in class_ranges:

        duplicate = False

        for prev in filtered:

            if (
                prev["from_force"] == r["from_force"] and
                prev["to_force"] == r["to_force"]
            ):
                duplicate = True
                break

        if not duplicate:
            filtered.append(r)

    class_ranges = filtered

        # ---------------------------------------
        # Remove duplicate ranges
        # ---------------------------------------

    # =====================================================
    # FORCE DISPLAY COLUMN ORDER
    # =====================================================
    if cert_type == "loadcell":

        display_columns = [

            "Applied Force kN",

            display_s1,
            display_s2,
            display_s3,
            display_s4,

            avg134_name,

            "REP_%",
            "REPR_%",
            "RESO_%",
            "ZERO_%",
            "CREEP_%",
            "INTERP_%",
            "W_INITIAL_%",
            "BMC_%",
            "FINAL_W_%",
            "ISO_CLASS"
        ]

        calc_df = calc_df[display_columns]

        calc_df.columns = [

            "Applied Force kN",

            "Position 0° series 1",
            "Position 0° series 2",
            "Position 120° series 3",
            "Position 240° series 4",

            "Average 1,3,4",

            "REP_%",
            "REPR_%",
            "RESO_%",
            "ZERO_%",
            "CREEP_%",
            "INTERP_%",
            "W_INITIAL_%",
            "BMC_%",
            "FINAL_W_%",
            "ISO_CLASS"
        ]

    else:

        base_cols = [

            "Applied Force kN",

            display_s1,
            display_s2,
            display_s3,
            display_s4,

            avg34_name,
            avg134_name
        ]

        remaining_cols = [

            c for c in calc_df.columns
            if c not in base_cols
        ]

        calc_df = calc_df[
            base_cols + remaining_cols
        ]


    # =====================================================
    # OVERALL CLASS
    # =====================================================

    overall_class = "Out of Class"

    if class_ranges:

        overall_class = (
            class_ranges[0]["class"]
            if class_ranges
            else "Out of Class"
        )

    return {

        "certificate_no": input_data["certificate_no"],
        "customer_name": input_data["customer_name"],
        "capacity": input_data["capacity"],
        "resolution": input_data["resolution"],
        "temperature": input_data["temperature"],
        "humidity": input_data["humidity"],

        "zero_error": zero_error,
        "creep": creep,
        "overall_class": overall_class,

        "force_to_avg_eq": force_to_avg_eq,
        "avg_to_force_eq": avg_to_force_eq,

        "table_df": calc_df,
        "export_df": export_df,
        "class_ranges": class_ranges,
    }


import re as _re

def strip_customer_prefix_for_filename(value):
    """
    Root cause of the filename bug: the filename is built from the raw
    customer name and then has "/" replaced with "-" (to make it a legal
    Windows filename). A customer name entered as "M/s. Roots Metrology"
    is therefore turned into "M-s. Roots Metrology" by that sanitization
    step, producing "...LC-20N-M-s. Roots Metrology...".

    This only strips a leading M/s./M/S./M/s/M/S prefix for the FILENAME.
    It must never be applied to the certificate content, which reproduces
    the customer name exactly as entered.
    """

    text = str(value or "").strip()

    return _re.sub(r"^(M/[sS]\.?\s*)", "", text, flags=_re.IGNORECASE)


def safe_value(value, default="Nil"):
    """
    Normalize a value that is about to be inserted into the certificate.

    Root cause of the "None" bug: Dash Input/State components return Python
    None (not "") when a field is left blank. That None was being passed
    straight into str(value), producing the literal text "None" in the
    generated certificate. Any blank, None, "nan" or "None" value is now
    rendered as "NIL" instead, everywhere a placeholder is filled in.
    """

    if value is None:
        return default

    text = str(value).strip()

    if text == "" or text.lower() in ("none", "nan"):
        return default

    return text


def replace_text_in_paragraph(paragraph, replacements):

    # ------------------------------------------------------------------
    # Pass 1: replace placeholders that live entirely inside a single run.
    # This is the normal case in these templates (verified against the
    # manual/template XML), and it is what preserves formatting correctly:
    # we only touch run.text, never run.bold / run.italic, so a bold label
    # run (e.g. "Customer Ref. No.") sitting next to a non-bold placeholder
    # run (e.g. ": {{CUSTOMER_REF_NO}}") keeps its own original formatting.
    # ------------------------------------------------------------------

    for run in paragraph.runs:

        if not run.text:
            continue

        new_text = run.text
        changed = False

        for key, value in replacements.items():

            if key in new_text:
                new_text = new_text.replace(key, safe_value(value))
                changed = True

        if changed:
            run.text = new_text

    # ------------------------------------------------------------------
    # Pass 2 (fallback): a placeholder that is still split across multiple
    # runs (e.g. Word inserted a spell-check boundary in the middle of a
    # "{{TAG}}") won't have been caught above. Collapse the paragraph text
    # into the first run in that case only. We deliberately do NOT force
    # bold/italic here anymore -- the previous code's forced
    # `bold = False` on runs[0] is exactly what was stripping bold off
    # labels like "Customer Ref. No." and "Accessories". Leaving the run's
    # existing formatting alone matches the manual certificate.
    # ------------------------------------------------------------------

    # Only the run(s) that make up the placeholder itself are rewritten.
    # The replacement text takes the formatting of the run in which the
    # placeholder starts, so runs before/after it (e.g. a bold label like
    # "Calibration:" sharing the paragraph with "{{TYPE_FORCE1}}" split
    # across runs) keep their own original formatting untouched -- unlike
    # the previous version, which dumped the merged text into runs[0] and
    # inherited whatever formatting runs[0] happened to have.
    made_change = True

    while made_change:

        made_change = False

        runs = paragraph.runs

        full_text = "".join(r.text for r in runs)

        match_key = None
        match_idx = -1

        for key in replacements:
            i = full_text.find(key)
            if i != -1:
                match_key = key
                match_idx = i
                break

        if match_key is None:
            break

        value = safe_value(replacements[match_key])

        start = match_idx
        end = match_idx + len(match_key)

        cursor = 0
        start_run = end_run = None
        start_off = end_off = 0

        for i, r in enumerate(runs):

            r_len = len(r.text)

            if start_run is None and cursor + r_len > start:
                start_run = i
                start_off = start - cursor

            if cursor + r_len >= end:
                end_run = i
                end_off = end - cursor
                break

            cursor += r_len

        if start_run is None or end_run is None:
            # Defensive guard only; should not happen since match_idx
            # was found in full_text built from these same runs.
            break

        before = runs[start_run].text[:start_off]
        after = runs[end_run].text[end_off:]

        if start_run == end_run:
            runs[start_run].text = before + value + after
        else:
            runs[start_run].text = before + value
            for i in range(start_run + 1, end_run):
                runs[i].text = ""
            runs[end_run].text = after

        made_change = True

def replace_placeholders_everywhere(doc, replacements):

        # normal paragraphs

        for paragraph in doc.paragraphs:

            replace_text_in_paragraph(
                paragraph,
                replacements
            )

        # tables

        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        replace_text_in_paragraph(
                            paragraph,
                            replacements
                        )

        # headers

        for section in doc.sections:

            header = section.header

            for paragraph in header.paragraphs:

                replace_text_in_paragraph(
                    paragraph,
                    replacements
                )

            for table in header.tables:

                for row in table.rows:

                    for cell in row.cells:

                        for paragraph in cell.paragraphs:

                            replace_text_in_paragraph(
                                paragraph,
                                replacements
                            )

            footer = section.footer

            for paragraph in footer.paragraphs:

                replace_text_in_paragraph(
                    paragraph,
                    replacements
                )

def replace_remarks(doc, remarks):

    def process_paragraph(paragraph):

        for run in paragraph.runs:

            if "{{REMARKS}}" in run.text:

                before = run.text.replace("{{REMARKS}}", "")
                run.text = before

                new_run = paragraph.add_run(str(remarks))
                new_run.bold = False
                new_run.italic = False

                return

    # Normal paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Headers & Footers
    for section in doc.sections:

        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph)

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)

        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph)

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)       
# =========================================================
# PROVING RING CERTIFICATE GENERATOR
# =========================================================

def generate_proving_ring_certificate(results, input_data):

        TEMPLATE_PATH = resource_path(
            os.path.join("templates", "C-Certificate no-Year-PR capacity, party name.docx")
        )

        doc = Document(TEMPLATE_PATH)

        cal_date_str = input_data.get("calibration_date")

        if cal_date_str:

            cal_date = datetime.strptime(
                cal_date_str,
                "%Y-%m-%d"
            )

        else:

            cal_date = datetime.today()

        expiry_date = cal_date + relativedelta(
            years=2,
            months=2
        )

        if input_data.get("test_type") == "Compression":

            type_force = "Compression Pads"
            type_force1 = "Compression"

        else:

            type_force = "Tension Shackles"
            type_force1 = "Tension"

        if input_data["certificate_type"] == "provingring":

            position_text = "180° and 360°"

        else:

            position_text = "120° and 240°"

        # ----------------------------------
        # ADD THESE LINES
        # ----------------------------------

        cert_type = input_data["certificate_type"]

        certificate_no = str(input_data.get("certificate_no", "")).strip()

        if not certificate_no.upper().startswith("C-"):
            certificate_no = f"C-{certificate_no}"

        pos3 = POSITION_MAP[cert_type]["s3"].replace(" series 3", "")
        pos4 = POSITION_MAP[cert_type]["s4"].replace(" series 4", "")

        # ----------------------------------

 


        replacements = {

           "{{CERTIFICATE_NO}}":
                certificate_no,
            
            "{{CASE_FILE}}":
                 input_data.get("case_file", ""),

            "{{TYPE_FORCE}}":
                type_force,    
             
            "{{TYPE_FORCE1}}":
                type_force1, 
                
            "{{POSITION_TEXT}}":
             position_text,
            
            "{{POS3}}":
             pos3,

            "{{POS4}}":
             pos4,

            "{{REMARKS}}": input_data["remarks"],
           
            "{{FORCE_MACHINE}}":
                input_data.get("force_machine", ""),

            "{{BMC}}":
                f"{input_data.get('bmc_value', 0):.3f}".rstrip("0").rstrip("."),

            "{{CUSTOMER_NAME}}":
                str(input_data.get("customer_name","")),

            "{{CUSTOMER_ADDRESS}}":
                str(input_data.get("customer_address","")),

            "{{CUSTOMER_REF_NO}}":
                str(input_data.get("customer_ref_no","")),

            "{{CUSTOMER_REF_DATE}}":
                str(input_data.get("customer_ref_date","")),

            "{{RING_TYPE}}":
                str(input_data.get("ring_type","")),

            "{{DIAL_NO}}":
                str(input_data.get("dial_no","")),

            "{{CAPACITY}}":
                str(input_data.get("capacity","")),

            "{{MANUFACTURER}}":
                str(input_data.get("manufacturer","")),

            "{{EQUIPMENT_SRNO}}":
                str(input_data.get("equipment_srno","")),

            "{{RESOLUTION}}":
                str(input_data.get("resolution","")),

            "{{SMALL_POINTER_SETTING}}":
                str(input_data.get(
                    "small_pointer_setting",
                    DEFAULT_SMALL_POINTER_SETTING
                )) or DEFAULT_SMALL_POINTER_SETTING,

            "{{CURRENT_DATE}}":
                datetime.today().strftime("%d.%m.%Y"),

            "{{CALIBRATION_DATE}}":
                cal_date.strftime("%d.%m.%Y"),

            "{{EXPIRY_DATE}}":
                expiry_date.strftime("%d.%m.%Y")
        }




        # =====================================================
        # REPLACE PLACEHOLDERS
        # =====================================================

        from copy import deepcopy

        for table in doc.tables:

            for row in table.rows:

                if "{{CLASSIFICATION_TABLE}}" in " ".join(
                    cell.text for cell in row.cells
                ):

                    template_row = row

                    for r in results["class_ranges"]:

                        new_row = table.add_row()

                        new_row.cells[0].text = f"Class {r['class']}"
                        new_row.cells[1].text = input_data["test_type"]
                        new_row.cells[2].text = f"{format_force(r['from_force'])} kN"
                        new_row.cells[3].text = f"{format_force(r['to_force'])} kN"
                        new_row.cells[4].text = r["uncertainty"]

                    table._tbl.remove(template_row._tr)

                    break

        replace_placeholders_everywhere(
                doc,
                replacements
        )
        
            

        # =====================================================
        # CALIBRATION TABLE
        # =====================================================

# =====================================================
# CALIBRATION TABLE
# =====================================================

        cert_type = input_data["certificate_type"]

        s1_col = POSITION_MAP[cert_type]["s1"]
        s2_col = POSITION_MAP[cert_type]["s2"]
        s3_col = POSITION_MAP[cert_type]["s3"]
        s4_col = POSITION_MAP[cert_type]["s4"]

        raw_df = pd.DataFrame(input_data["table_data"])

# Use processed results table instead of raw table

        df = results["export_df"].copy()

        for table in doc.tables:

            try:

                placeholder_row_idx = None

                for i, row in enumerate(table.rows):

                    row_text = " ".join(
                        cell.text for cell in row.cells
                    )

                    if "{{DATA_ROW}}" in row_text:

                        placeholder_row_idx = i
                        break

                if placeholder_row_idx is None:
                    continue

                from copy import deepcopy

                template_row = table.rows[placeholder_row_idx]

                for _, data_row in df.iterrows():

                    new_tr = deepcopy(template_row._tr)

                    table._tbl.append(new_tr)

                    new_row = table.rows[-1]

                    values = data_row.tolist()

                    for col_idx, value in enumerate(values):

                        if pd.isna(value):
                            value = ""

                        if col_idx < len(new_row.cells):
                            new_row.cells[col_idx].text = str(value)

                table._tbl.remove(template_row._tr)

            except Exception as e:

                print("TABLE ERROR:", e)

            except:
                pass
        # =====================================================
        # FILE NAME
        # =====================================================

        filename = (
            f"{certificate_no}"
            f"-{datetime.today().year}"
            f"-PR-"
            f"{input_data.get('capacity','')}kN-"
            f"{safe_value(strip_customer_prefix_for_filename(input_data.get('customer_name','')))}.docx"
        )

        filename = filename.replace("/", "-")
        filename = filename.replace("\\", "-")

        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)

        print(f"Certificate Saved : {filepath}")
        logging.info("PR certificate saved: %s", filepath)

        return filename



def generate_loadcell_certificate(results, input_data):

    TEMPLATE_PATH = resource_path(
        os.path.join("templates", "C-Certificate no-Year-LC capacity, party name.docx")
    )

    doc = Document(TEMPLATE_PATH)

    cal_date_str = input_data.get("calibration_date")

    if cal_date_str:
        cal_date = datetime.strptime(
            cal_date_str,
            "%Y-%m-%d"
        )
    else:
        cal_date = datetime.today()

    expiry_date = cal_date + relativedelta(
        years=2,
        months=2
    )

    if input_data.get("test_type") == "Compression":

        type_force = "Compression Pads"
        type_force1 = "Compression"

    else:

        type_force = "Tension Shackles"
        type_force1 = "Tension"

    # Preserve the customer name exactly as the user typed it. The template
    # no longer contains a hardcoded "M/s" prefix, so whatever the user
    # enters (including "M/s", "M/S.", etc., or nothing at all) must be
    # reproduced verbatim -- it is not our job to add, remove, or normalize
    # any prefix.
    customer_name_clean = str(input_data.get("customer_name", "")).strip()

    certificate_no = str(input_data.get("certificate_no", "")).strip()

    if not certificate_no.upper().startswith("C-"):
        certificate_no = f"C-{certificate_no}"

    replacements = {

       "{{CERTIFICATE_NO}}":
            certificate_no, 

        "{{CASE_FILE}}":
            str(input_data.get("case_file","")),

       "{{CUSTOMER_NAME}}":
            customer_name_clean,

        "{{CUSTOMER_ADDRESS}}":
            str(input_data.get("customer_address","")),

        "{{CUSTOMER_REF_NO}}":
            str(input_data.get("customer_ref_no","")),

        "{{CUSTOMER_REF_DATE}}":
            str(input_data.get("customer_ref_date","")),

        "{{CAPACITY}}":
            str(input_data.get("capacity","")),

        "{{MANUFACTURER}}":
            str(input_data.get("manufacturer","")),

        "{{MAKE}}":
            str(input_data.get("make","")),

        "{{EQUIPMENT_SRNO}}":
            str(input_data.get("equipment_srno","")),

        "{{RESOLUTION}}":
            str(input_data.get("resolution","")),

        # Resolution UNIT (div / kN / count), selected via the Resolution
        # dropdown. Used both in the instrument description table's
        # "Resolution: {{RESOLUTION}} {{RESOLUTION_UNIT}}" row and in the
        # Interpolation Equation's "X = Indicator reading in ..." line.
        # Single source of truth: input_data["resolution_unit"] (never
        # hard-coded here) -- kept identical to {{NO_LOAD_OUTPUT_UNIT}}.
        "{{RESOLUTION_UNIT}}":
            str(input_data.get(
                "resolution_unit",
                DEFAULT_RESOLUTION_UNIT
            )) or DEFAULT_RESOLUTION_UNIT,

        "{{FORCE_MACHINE}}":
            input_data.get("force_machine", ""),

        "{{REMARKS}}": input_data["remarks"],

        "{{BMC}}":
            f"{input_data.get('bmc_value', 0):.3f}".rstrip("0").rstrip("."),

        "{{DIGITAL_INDICATOR_SRNO}}":
            str(input_data.get("digital_indicator_srno","")),

        "{{CONNECTOR_TYPE}}":
            str(input_data.get("connector_type","")),

        "{{CABLE_LENGTH}}":
            str(input_data.get("cable_length","")),

        "{{MODEL_NO}}":
            str(input_data.get("indicator_model_no","")),

        "{{NO_LOAD_OUTPUT}}":
            str(input_data.get(
                "no_load_output",
                DEFAULT_NO_LOAD_OUTPUT
            )) or DEFAULT_NO_LOAD_OUTPUT,

        # Resolution UNIT (div / kN / count) selected via the Resolution
        # dropdown -- drives the "(Digital Indicator Reading in ...)" text
        # in the Results section. This is intentionally separate from the
        # {{RESOLUTION}} placeholder above, which is the numeric Resolution
        # VALUE (e.g. 0.001) and must never be overwritten by the unit.
        "{{NO_LOAD_OUTPUT_UNIT}}":
            str(input_data.get(
                "resolution_unit",
                DEFAULT_RESOLUTION_UNIT
            )) or DEFAULT_RESOLUTION_UNIT,

        "{{TYPE_FORCE}}":
            type_force,

        "{{TYPE_FORCE1}}":
            type_force1,

        "{{AVG_TO_FORCE_EQ}}":
                results.get("avg_to_force_eq",""),

        "{{FORCE_TO_AVG_EQ}}":
                results.get("force_to_avg_eq",""),       

        "{{CURRENT_DATE}}":
            datetime.today().strftime("%d.%m.%Y"),

        "{{CALIBRATION_DATE}}":
            cal_date.strftime("%d.%m.%Y"),

        "{{EXPIRY_DATE}}":
            expiry_date.strftime("%d.%m.%Y"),


        "{{RING_TYPE}}": "Load Cell",

        "{{POSITION_TEXT}}": "120° and 240°",

        "{{POS3}}": "Position 120°",

        "{{POS4}}": "Position 240°",
    }

    for table in doc.tables:

        for row in table.rows:

            if "{{CLASSIFICATION_TABLE}}" in " ".join(
                cell.text for cell in row.cells
            ):

                template_row = row

                for r in results["class_ranges"]:

                    new_row = table.add_row()

                    new_row.cells[0].text = f"Class {r['class']}"
                    new_row.cells[1].text = input_data["test_type"]
                    new_row.cells[2].text = f"{format_force(r['from_force'])} kN"
                    new_row.cells[3].text = f"{format_force(r['to_force'])} kN"
                    new_row.cells[4].text = r["uncertainty"]

                table._tbl.remove(template_row._tr)

                break


    replace_placeholders_everywhere(
        doc,
        replacements
    )


    df = results["export_df"].copy()

    for table in doc.tables:

        try:

            placeholder_row_idx = None

            for i, row in enumerate(table.rows):

                row_text = " ".join(
                    cell.text for cell in row.cells
                )

                if "{{DATA_ROW}}" in row_text:

                    placeholder_row_idx = i
                    break

            if placeholder_row_idx is None:
                continue

            from copy import deepcopy

            template_row = table.rows[placeholder_row_idx]

            for _, data_row in df.iterrows():

                new_tr = deepcopy(template_row._tr)

                table._tbl.append(new_tr)

                new_row = table.rows[-1]

                values = data_row.tolist()

                for col_idx, value in enumerate(values):

                    if pd.isna(value):
                        value = ""

                    if col_idx < len(new_row.cells):
                        new_row.cells[col_idx].text = str(value)

            table._tbl.remove(template_row._tr)

        except Exception as e:

            print("TABLE ERROR:", e)

    filename = (
        f"{certificate_no}"
        f"-{datetime.today().year}"
        f"-LC-"
        f"{input_data.get('capacity','')}-"
        f"{safe_value(strip_customer_prefix_for_filename(input_data.get('customer_name','')))}.docx"
    )

    filename = filename.replace("/", "-")
    filename = filename.replace("\\", "-")

    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    print(f"Certificate Saved : {filepath}")
    logging.info("LC certificate saved: %s", filepath)

    return filename   
# =========================================================
# APP LAYOUT
# =========================================================

print("=" * 60)
print("USING UPDATED main.final.py")
print("=" * 60)

app.layout = dbc.Container([

    # =====================================================
    # HEADER
    # =====================================================

    html.Div(

        [

            # Left Logo
            html.Div(

                html.Img(
                    src="/assets/csir_logo.png",
                    style={
                        "height": "120px",
                        "width": "120px",
                        "objectFit": "contain",
                        "backgroundColor": "white",
                        "padding": "10px",
                        "borderRadius": "10px"
                    }
                ),

                style={
                    "display": "flex",
                    "alignItems": "center",
                    "justifyContent": "center",
                    "marginRight": "25px"
                }
            ),

            # Right Text
            html.Div(

                [

                    html.H1(
                        "DIGIForce-CCS™: Digital Force Calibration, Classification & Certification Software",
                        style={
                            "color": "white",
                            "fontWeight": "700",
                            "fontSize": "48px",
                            "marginBottom": "10px"
                        }
                    ),

                    html.Hr(
                        style={
                            "borderTop": "2px solid #66c2ff",
                            "margin": "8px 0"
                        }
                    ),

                    html.H4(
                        "Developed By CSIR-NPL Force & Hardness Laboratory",
                        style={
                            "color": "white",
                            "fontWeight": "500"
                        }
                    )

                ],

                style={
                    "flex": "1"
                }

            )

        ],

        style={
            "display": "flex",
            "alignItems": "center",
            "background": "linear-gradient(90deg,#001B3A,#004B8D)",
            "padding": "25px",
            "borderRadius": "0px 0px 20px 20px",
            "marginBottom": "20px",
            "boxShadow": "0 4px 15px rgba(0,0,0,0.3)"
        }

    ),

    # =====================================================
    # CERTIFICATE DETAILS
    # =====================================================

    dbc.Card([

        dbc.CardHeader(
            "Certificate Details",
            style=CARD_HEADER_STYLE
        ),

        dbc.CardBody([

            dbc.Row([
                dbc.Col([
                    dbc.Label("Case File Number"),
                    dbc.Input(
                        id="case_file",
                        type="text"
                    )
                ])
            ]),

            dbc.Row([

                dbc.Col([

                    html.Label("Certificate Number", style=LABEL_STYLE),

                    dcc.Input(
                        id="certificate_no",
                        className="form-control",
                        style=INPUT_STYLE
                    )

                ], md=3),

                dbc.Col([

                    html.Label("Calibration Date", style=LABEL_STYLE),

                    dcc.DatePickerSingle(
                        id="calibration_date"
                    )

                ], md=3),

                dbc.Col([

                    html.Label("Description of Equipment", style=LABEL_STYLE),

                    dcc.Dropdown(
                        id="certificate_type",
                        options=[
                            {
                                "label": "Load Cell",
                                "value": "loadcell"
                            },
                            {
                                "label": "Proving Ring",
                                "value": "provingring"
                            }
                        ],
                        value="loadcell"
                    )

                ], md=3),

                dbc.Col([

                    html.Label("Type of Test Performed", style=LABEL_STYLE),

                    dcc.Dropdown(
                        id="test_type",
                        options=[
                            {
                                "label": "Compression",
                                "value": "Compression"
                            },
                            {
                                "label": "Tension",
                                "value": "Tension"
                            }
                        ],
                        value="Compression"
                    )

                ], md=3)

            ])

        ])

    ], style=CARD_STYLE),

    # =====================================================
    # CUSTOMER DETAILS
    # =====================================================

    dbc.Card([

        dbc.CardHeader(
            "Customer Details",
            style=CARD_HEADER_STYLE
        ),

        dbc.CardBody([

            dbc.Row([

                dbc.Col([
                    html.Label("Customer Name", style=LABEL_STYLE),
                    dcc.Input(
                        id="customer_name",
                        type="text",

                        className="form-control",
                        style={
                                    "width": "100%",
                                    "fontSize": "14px",
                                    "height": "42px"
                            }
                    )
                ], md=4),

            

                dbc.Col([
                    html.Label("Customer Reference No.", style=LABEL_STYLE),
                    dcc.Input(
                        id="customer_ref_no",
                        className="form-control",
                        style=INPUT_STYLE
                    )
                ], md=4),

                dbc.Col([
                    html.Label("Reference Date", style=LABEL_STYLE),
                    dcc.DatePickerSingle(
                        id="customer_ref_date"
                    )
                ], md=4)

            ]),

            html.Br(),

            dbc.Row([

                dbc.Col([
                    html.Label("Customer Address", style=LABEL_STYLE),
                    dcc.Textarea(
                        id="customer_address",
                        className="form-control"
                    )
                ])

            ])

        ])

    ], style=CARD_STYLE),

    # =====================================================
    # INSTRUMENT DETAILS
    # =====================================================

    dbc.Card([

        dbc.CardHeader(
            "Instrument Details",
            style=CARD_HEADER_STYLE
        ),

        dbc.CardBody([
            html.Div(id="dynamic_instrument_fields")
        ])

    ], style=CARD_STYLE),

    # =====================================================
    # ENVIRONMENTAL CONDITIONS
    # =====================================================

    dbc.Card([

        dbc.CardHeader(
            "Environmental Conditions",
            style=CARD_HEADER_STYLE
        ),

        dbc.CardBody([

            dbc.Row([

                dbc.Col([
                    html.Label("Temperature (°C)", style=LABEL_STYLE),
                    dcc.Input(
                        id="temperature",
                        type="number",
                        className="form-control",
                        style=INPUT_STYLE
                    )
                ], md=4),

                dbc.Col([
                    html.Label("Humidity (%)", style=LABEL_STYLE),
                    dcc.Input(
                        id="humidity",
                        type="number",
                        className="form-control",
                        style=INPUT_STYLE
                    )
                ], md=4),
                    
            


            ]),



        ])

    ], style=CARD_STYLE),

   
    # =====================================================
    # CALIBRATION TABLE
    # =====================================================

    dbc.Card([

        dbc.CardHeader(
            "Calibration Readings",
            style=CARD_HEADER_STYLE
        ),

        dbc.CardBody([

            dbc.Row([

                dbc.Col([

                    html.Label("Force Standard Machine", style=LABEL_STYLE),

                    dcc.Dropdown(

                        id="force_machine",

                        options=[

                            {"label":"1 kN Dead Weight Force Machine","value":"1"},
                            {"label":"5 kN Dead Weight Force Machine","value":"2"},
                            {"label":"50 kN Dead Weight Force Machine","value":"3"},
                            {"label":"100 kN Dead Weight Force Machine","value":"4"},
                            {"label":"1 MN FSM Force Machine","value":"5"},
                            {"label":"1000 kN Lever Multiplication Force Machine","value":"6"},
                            {"label":"1000 kN HMS Force Machine","value":"7"},
                            {"label":"200 kN Force Machine","value":"8"},
                            {"label":"3000 kN Force Machine","value":"9"}

                        ],

                        value="4"

                    )

                ], md=4),

                dbc.Col([
                    html.Div(
                        build_reading_setting_field("loadcell"),
                        id="dynamic_reading_setting_field"
                    )
                ], md=4),

            ]),
            html.Br(),

            dash_table.DataTable(

                id="calibration_table",

                columns=[
                    {
                        "name": i,
                        "id": i,
                        "editable": True
                    }
                    for i in default_df.columns
                ],

                data=default_df.to_dict("records"),

                editable=True,
                row_deletable=True,

                style_table={
                    "overflowX": "auto"
                },

                style_header={
                    "backgroundColor": "#003F73",
                    "color": "white",
                    "fontWeight": "600"
                },

                style_cell={
                    "textAlign": "center",
                    "padding": "8px"
                }

            ),

            html.Br(),

            dbc.Button(
                "Add Row",
                id="add_row",
                color="primary",
                style=BUTTON_STYLE
            ),

            html.Hr(style={"margin": "15px 0"}),

            html.Div(
                "Import data from Excel",
                style={
                    "fontWeight": "600",
                    "marginBottom": "8px"
                }
            ),

            dcc.Upload(
                id="upload_excel",
                children=dbc.Button(
                    "Upload Excel Template",
                    color="info",
                    style=BUTTON_STYLE
                ),
                multiple=False
            ),

            html.Div(
                id="excel_upload_status",
                style={
                    "marginTop": "8px",
                    "fontWeight": "600"
                }
            )

        ])

    ], style=CARD_STYLE),

    html.Br(),

    dbc.Row([

        dbc.Col(
            dbc.Button(
                "Compute",
                id="compute_btn",
                color="warning",
                style=BUTTON_STYLE
            ),
            width="auto"
        ),

        dbc.Col(
            dbc.Button(
                "Clear Result",
                id="clear_result_btn",
                color="secondary",
                style=BUTTON_STYLE
            ),
            width="auto",
            className="ms-auto"
        ),

    ], justify="between", className="g-0"),

    html.Br(),
    html.Br(),



    html.Div(
        id="classification_output"
    ),

    html.Br(),

    html.Div(

        id="remarks_section",

        style={"display":"none"},

        children=[

            dbc.Card(

                [

                    dbc.CardHeader(
                        "Remarks",
                        style=CARD_HEADER_STYLE
                    ),

                    dbc.CardBody(

                        [

                            dcc.Textarea(

                                id="remarks_box",

                                style={

                                    "width":"100%",

                                    "height":"140px"

                                }

                            )

                        ]

                    )

                ],

                style=CARD_STYLE

            ),

        ]

    ),

    html.Br(),

    dbc.Button(

        "Generate Certificate",

        id="generate_certificate_btn",

        color="success",

        disabled=True,

        style=BUTTON_STYLE

    ),

    html.Br(),

    html.Div(id="certificate_status"),

    html.Br(),


], fluid=True)


# =========================================================
# DYNAMIC EQUIPMENT FORM
# =========================================================

@app.callback(
    Output("dynamic_instrument_fields", "children"),
    Input("certificate_type", "value")
)

def update_instrument_fields(cert_type):

    if cert_type == "loadcell":

        return html.Div([

            dbc.Row([

                dbc.Col([
                    html.Label("Digital Indicator Sr. No.", style=LABEL_STYLE),
                    dcc.Input(id="digital_indicator_srno", className="form-control")
                ], md=3),

                dbc.Col([

                    html.Label("Capacity", style=LABEL_STYLE),

                    dbc.InputGroup([

                        dcc.Input(
                            id="capacity",
                            type="number",
                            className="form-control"
                        ),

                        dcc.Dropdown(
                            id="capacity_unit",
                            options=[
                                {"label": "N", "value": "N"},
                                {"label": "kN", "value": "kN"}
                            ],
                            value="kN",
                            style={"width": "120px"}
                        )

                    ])

                ], md=3),

                
                dbc.Col([
                    html.Label("Manufacturer", style=LABEL_STYLE),
                    dcc.Input(id="manufacturer", className="form-control")
                ], md=3),


                dbc.Col([
                    html.Label("Equipment S.No.", style=LABEL_STYLE),
                    dcc.Input(id="equipment_srno", className="form-control")
                ], md=3)

            ]),

            html.Br(),

            dbc.Row([

                dbc.Col([
                    html.Label("Model No.", style=LABEL_STYLE),
                    dcc.Input(id="indicator_model_no", className="form-control")
                ], md=3),

                dbc.Col([
                    html.Label("Connector Type", style=LABEL_STYLE),
                    dcc.Input(id="connector_type", className="form-control")
                ], md=3),

                dbc.Col([

                    html.Label("Resolution", style=LABEL_STYLE),

                    dbc.InputGroup([

                        dcc.Input(
                            id="resolution",
                            className="form-control"
                        ),

                        dcc.Dropdown(
                            id="resolution_unit",
                            options=RESOLUTION_UNIT_OPTIONS,
                            value=DEFAULT_RESOLUTION_UNIT,
                            clearable=False,
                            style={"width": "120px"}
                        )

                    ])

                ], md=3),

                dbc.Col([
                    html.Label("Cable Length", style=LABEL_STYLE),
                    dcc.Input(id="cable_length",   type="number", className="form-control")
                ], md=3)

            ]),

            html.Br(),

            dbc.Row([

                dbc.Col([
                    html.Label("Make", style=LABEL_STYLE),
                    dcc.Input(id="make", className="form-control")
                ], md=3)

            ]),


        ])

    return html.Div([

        dbc.Row([

            dbc.Col([

                html.Label("Type", style=LABEL_STYLE),

                dcc.Dropdown(
                    id="ring_type",
                    options=[
                        {
                            "label": "Integral Ring",
                            "value": "Integral Ring"
                        },
                        {
                            "label": "Bow Dynamometer",
                            "value": "Bow Dynamometer"
                        }
                    ],
                    value="Integral Ring"
                )

            ], md=3),

            dbc.Col([
                html.Label("Dial Gauge Sr. No.", style=LABEL_STYLE),
                dcc.Input(id="dial_no", className="form-control")
            ], md=3),

            dbc.Col([

                html.Label("Capacity", style=LABEL_STYLE),

                dbc.InputGroup([

                    dcc.Input(
                        id="capacity",
                        className="form-control"
                    ),

                    dcc.Dropdown(
                        id="capacity_unit",
                        options=[
                            {"label": "N", "value": "N"},
                            {"label": "kN", "value": "kN"}
                        ],
                        value="kN",
                        style={"width": "120px"}
                    )

                ])

            ], md=3),

            dbc.Col([
                html.Label("Manufacturer", style=LABEL_STYLE),
                dcc.Input(id="manufacturer", className="form-control")
            ], md=3)

        ]),

        html.Br(),

        dbc.Row([

            dbc.Col([
                html.Label("Sl. No.", style=LABEL_STYLE),
                dcc.Input(id="equipment_srno", className="form-control")
            ], md=3),

            dbc.Col([

                html.Label("Resolution", style=LABEL_STYLE),

                dbc.InputGroup([

                    dcc.Input(
                        id="resolution",
                        className="form-control"
                    ),

                    dcc.Dropdown(
                        id="resolution_unit",
                        options=RESOLUTION_UNIT_OPTIONS,
                        value=DEFAULT_RESOLUTION_UNIT,
                        clearable=False,
                        style={"width": "120px"}
                    )

                ])

            ], md=3),

            dbc.Col([

                html.Label("1 Div. Value", style=LABEL_STYLE),

                dbc.InputGroup([

                    dcc.Input(
                    id="div_value",
                    type="number",
                    className="form-control"
                    ),

                    dcc.Dropdown(
                        id="div_unit",
                        options=[
                            {
                                "label": "m",
                                "value": "m"
                            },
                            {
                                "label": "mm",
                                "value": "mm"
                            }
                        ],
                        value="mm",
                        style={"width": "120px"}
                    )

                ])

            ], md=6)

        ])

    ])

# =========================================================
# DYNAMIC READING-SETTING FIELD
# (Small Pointer Setting for Proving Ring / No Load Output for Load Cell)
# =========================================================

@app.callback(
    Output("dynamic_reading_setting_field", "children"),
    Input("certificate_type", "value")
)
def update_reading_setting_field(cert_type):
    print("========== CALLBACK ==========")
    print("Certificate Type:", cert_type)
    print("==============================")
    return build_reading_setting_field(cert_type)


# =========================================================
# UPDATE COLUMN NAMES
# =========================================================
# =========================================================
# UPDATE COLUMN NAMES AUTOMATICALLY
# =========================================================
@app.callback(
    Output("calibration_table", "columns"),
    Output("calibration_table", "data", allow_duplicate=True),
    Input("certificate_type", "value"),
    prevent_initial_call=True
)
def update_columns(cert_type):

    pos3 = POSITION_MAP[cert_type]["s3"]
    pos4 = POSITION_MAP[cert_type]["s4"]

    cols = [

        "Applied Force kN",

        "Position 0° series 1",

        "Position 0° series 2",

        pos3,

        pos4,

        "Average 1,3,4"
    ]

    columns = [
        {
            "name": c,
            "id": c,
            "editable": c != "Average 1,3,4"
        }
        for c in cols
    ]

    data = [
        {c: "" for c in cols}
    ]

    return columns, data
# =========================================================
# ADD ROW
# =========================================================

@app.callback(
    Output("calibration_table", "data", allow_duplicate=True),
    Input("add_row", "n_clicks"),
    State("calibration_table", "data"),
    State("calibration_table", "columns"),
    prevent_initial_call=True
)

def add_row(n_clicks, rows, columns):

    rows.append({
        c["id"]: ""
        for c in columns
    })

    return rows


# =========================================================
# AUTO UPDATE AVERAGE
# =========================================================

@app.callback(
    Output("calibration_table", "data", allow_duplicate=True),
    Input("calibration_table", "data_timestamp"),
    State("calibration_table", "data"),
    State("certificate_type", "value"),
    State("resolution", "value"),
    prevent_initial_call=True
)

def update_average(timestamp, rows, cert_type, resolution):

    if rows is None:
        return rows

    # =====================================================
    # DETERMINE COLUMN NAMES
    # =====================================================

    s3_col = POSITION_MAP[cert_type]["s3"]
    s4_col = POSITION_MAP[cert_type]["s4"]

    updated_rows = []

    for row in rows:

        try:

            s1 = float(
                row.get("Position 0° series 1", 0) or 0
            )

            s3 = float(
                row.get(s3_col, 0) or 0
            )

            s4 = float(
                row.get(s4_col, 0) or 0
            )

            
            try:

                values = []

                if row.get("Position 0° series 1") not in [None, ""]:
                    values.append(float(row["Position 0° series 1"]))

                if row.get(s3_col) not in [None, ""]:
                    values.append(float(row[s3_col]))

                if row.get(s4_col) not in [None, ""]:
                    values.append(float(row[s4_col]))

                if len(values) > 0:

                    reading_decimals = get_decimal_places_from_resolution(
                        resolution
                    )

                    if reading_decimals is None:

                        # Fallback: same behaviour as before, if
                        # Resolution hasn't been entered yet.
                        s1_text = str(row.get("Position 0° series 1", "")).strip()

                        if "." in s1_text:
                            reading_decimals = len(s1_text.split(".")[1])
                        else:
                            reading_decimals = 0

                    avg_value = sum(values) / len(values)

                    row["Average 1,3,4"] = (
                        f"{avg_value:.{reading_decimals}f}"
                    )

                else:
                    row["Average 1,3,4"] = ""

            except Exception:
                row["Average 1,3,4"] = ""

        except Exception:

            row["Average 1,3,4"] = ""

        updated_rows.append(row)
        print("UPDATE AVERAGE CALLBACK TRIGGERED")
        print(rows[0] if rows else "NO ROWS")
    return updated_rows
# =========================================================
# TEMPORARY BACKEND LINK TEST
# =========================================================

@app.callback(
    Output("calibration_table", "data", allow_duplicate=True),
    Output("excel_upload_status", "children"),
    Input("upload_excel", "contents"),
    State("upload_excel", "filename"),
    State("certificate_type", "value"),
    prevent_initial_call=True
)
def upload_excel_data(contents, filename, cert_type):

    if contents is None:
        return dash.no_update, ""

    try:

        # ------------------------------------------------
        # Check file
        # ------------------------------------------------
        if not filename.lower().endswith((".xlsx", ".xls")):

            return (
                dash.no_update,
                "❌ Please upload an Excel file (.xlsx or .xls)."
            )

        # ------------------------------------------------
        # Decode uploaded Excel file
        # ------------------------------------------------
        content_type, content_string = contents.split(",")

        decoded = base64.b64decode(content_string)

        excel_file = io.BytesIO(decoded)

        # ------------------------------------------------
        # Read the common template
        # ------------------------------------------------
        df = pd.read_excel(
            excel_file,
            header=0
        )

        # Clean headers
        df.columns = [
            str(col).strip()
            for col in df.columns
        ]

        # ------------------------------------------------
        # Common template columns
        # ------------------------------------------------
        required_columns = [
            "Force kN",
            "Series 1",
            "Series 2",
            "Series 3",
            "Series 4"
        ]

        missing = [
            col for col in required_columns
            if col not in df.columns
        ]

        if missing:

            return (
                dash.no_update,
                "❌ Invalid template. Missing: "
                + ", ".join(missing)
            )

        # Keep only template data
        df = df[required_columns].copy()

        # ------------------------------------------------
        # Remove completely blank rows
        # ------------------------------------------------
        df = df.dropna(
            how="all"
        )

        # ------------------------------------------------
        # Convert to numeric
        # ------------------------------------------------
        for col in required_columns:

            df[col] = pd.to_numeric(
                df[col],
                errors="coerce"
            )

        # ------------------------------------------------
        # Remove rows where Force is blank
        # ------------------------------------------------
        df = df.dropna(
            subset=["Force kN"]
        )

        # =================================================
        # MAP TEMPLATE → EXISTING BACKGROUND COLUMN NAMES
        # =================================================

        if cert_type == "loadcell":

            df.rename(
                columns={
                    "Force kN":
                        "Applied Force kN",

                    "Series 1":
                        "Position 0° series 1",

                    "Series 2":
                        "Position 0° series 2",

                    "Series 3":
                        "Position 120° series 3",

                    "Series 4":
                        "Position 240° series 4"
                },
                inplace=True
            )

        elif cert_type == "provingring":

            df.rename(
                columns={
                    "Force kN":
                        "Applied Force kN",

                    "Series 1":
                        "Position 0° series 1",

                    "Series 2":
                        "Position 0° series 2",

                    "Series 3":
                        "Position 180° series 3",

                    "Series 4":
                        "Position 360° series 4"
                },
                inplace=True
            )

        else:

            return (
                dash.no_update,
                "❌ Please select equipment type first."
            )

        # ------------------------------------------------
        # Calculate Average 1,3,4 immediately after upload
        # ------------------------------------------------

        s1 = "Position 0° series 1"

        if cert_type == "loadcell":
            s3 = "Position 120° series 3"
            s4 = "Position 240° series 4"
        else:
            s3 = "Position 180° series 3"
            s4 = "Position 360° series 4"

        df["Average 1,3,4"] = (
            pd.to_numeric(df[s1], errors="coerce")
            + pd.to_numeric(df[s3], errors="coerce")
            + pd.to_numeric(df[s4], errors="coerce")
        ) / 3

        # Keep the same precision as your existing table
        df["Average 1,3,4"] = df["Average 1,3,4"].round(6)

        # Convert to table records
        data = df.to_dict("records")

        return (
            data,
            f"✓ {filename} loaded successfully "
            f"({len(data)} rows)"
        )

    except Exception as e:

        print(
            "Excel upload error:",
            e
        )

        return (
            dash.no_update,
            f"❌ Excel upload failed: {str(e)}"
        )
# =========================================================
# BACKEND CONNECTION TEST
# =========================================================

@app.callback(
    Output("classification_output","children"),


    Output("remarks_box","value"),

    Output("generate_certificate_btn","disabled"),
    Output("certificate_status","children"),
    Output("remarks_section", "style"),

    Input("compute_btn", "n_clicks"),

    Input("generate_certificate_btn","n_clicks"),
    
    State("calibration_table", "data"),

    State("certificate_no", "value"),
    State("case_file", "value"),
    State("test_type", "value"),
    State("customer_name", "value"),
    State("customer_address", "value"),
    State("customer_ref_no", "value"),
    State("customer_ref_date", "date"),

    State("capacity", "value"),
    State("capacity_unit", "value", allow_optional=True),
    State("resolution", "value"),
    State("resolution_unit", "value", allow_optional=True),

    State("ring_type", "value", allow_optional=True),
    State("dial_no", "value", allow_optional=True),
    State("manufacturer", "value", allow_optional=True),
    State("make", "value", allow_optional=True),
    State("equipment_srno", "value", allow_optional=True),

    State("digital_indicator_srno", "value", allow_optional=True),
    State("connector_type", "value", allow_optional=True),
    State("cable_length", "value", allow_optional=True),
    State("indicator_model_no", "value", allow_optional=True),


    State("calibration_date", "date"),

    State("div_value", "value", allow_optional=True),

    State("temperature", "value"),
    State("humidity", "value"),
    State("force_machine","value"),
    State("certificate_type", "value"),
    State("remarks_box","value"),

    State("small_pointer_setting", "value", allow_optional=True),
    State("no_load_output", "value", allow_optional=True),
    prevent_initial_call=True
)

def compute_classification(
        n_clicks,
        generate_clicks,

        table_data,

        certificate_no,
        case_file,
        test_type,
        customer_name,
        customer_address,
        customer_ref_no,
        customer_ref_date,

        capacity,
        capacity_unit,
        resolution,
        resolution_unit,

        ring_type,
        dial_no,
        manufacturer,
        make,
        equipment_srno,

        digital_indicator_srno,
        connector_type,
        cable_length,
        indicator_model_no,

        calibration_date,

        div_value,

        temperature,
        humidity,
        force_machine,
        certificate_type,
        remarks_box,

        small_pointer_setting,
        no_load_output):
    
    trigger = ctx.triggered_id
    print("\n===== RAW TABLE RECEIVED =====")

    for r in table_data:
        print(r)

    print("=============================\n")

    clean_rows = []

    clean_rows = []

    for row in table_data:
        clean_rows.append(row.copy())

    input_data = {

        "certificate_no": certificate_no,
        "case_file": case_file,
        "test_type": test_type,
        "customer_name": customer_name,
        "customer_address": customer_address,
        "customer_ref_no": customer_ref_no,
        "customer_ref_date": customer_ref_date,

        "capacity": f"{capacity} {capacity_unit}",
        "resolution": resolution,
        "resolution_unit": resolution_unit or DEFAULT_RESOLUTION_UNIT,

        "ring_type": ring_type,
        "dial_no": dial_no,
        "manufacturer": manufacturer,
        "make": make,
        "equipment_srno": equipment_srno,

        "digital_indicator_srno": digital_indicator_srno,
        "connector_type": connector_type,
        "cable_length": (
            cable_length
            if certificate_type == "loadcell"
            else ""
        ),

        "div_value": (
            div_value
            if certificate_type == "provingring"
            else ""
        ),
        "indicator_model_no": indicator_model_no,

        "calibration_date": calibration_date,

        "temperature": temperature,
        "humidity": humidity,
        "force_machine": FORCE_MACHINE_MAP[force_machine],

        "bmc_value": BMC_VALUES[force_machine],
        "table_data": clean_rows,

        "certificate_type": certificate_type,

        "small_pointer_setting": (
            (small_pointer_setting or DEFAULT_SMALL_POINTER_SETTING).strip()
            if certificate_type == "provingring"
            else ""
        ),

        "no_load_output": (
            str(no_load_output).strip()
            if certificate_type == "loadcell" and no_load_output not in (None, "")
            else (DEFAULT_NO_LOAD_OUTPUT if certificate_type == "loadcell" else "")
        ),
    }




# ==========================================================
# NUMERIC VALIDATION
# ==========================================================

    numeric_fields = {}

    # Required for both Load Cell and Proving Ring
    numeric_fields["Capacity"] = capacity
    numeric_fields["Resolution"] = resolution
    numeric_fields["Temperature"] = temperature
    numeric_fields["Humidity"] = humidity

    if certificate_type == "loadcell":

        # Required only for Load Cell
        numeric_fields["Cable Length"] = cable_length

    elif certificate_type == "provingring":

        # Required only for Proving Ring
        numeric_fields["1 Div. Value"] = div_value

    for field, value in numeric_fields.items():

        if value in [None, ""]:

                return (
                    dbc.Alert(
                        f"Wrong values entered. '{field}' must be numeric.",
                        color="danger"
                    ),
                    dash.no_update,
                    True,
                    "",
                    {"display": "none"}

                )

        try:
            float(value)

        except (TypeError, ValueError):

                return (
                    dbc.Alert(
                        f"Wrong values entered. '{field}' must be numeric.",
                        color="danger"
                    ),
                    dash.no_update,
                    True,
                    "",
                    {"display": "none"}

                )

    print("========== DEBUG ==========")
    print(type(input_data))
    print(input_data)
    print("===========================")

    results = process_calibration_data(input_data)


    if certificate_type=="loadcell":

        remarks = (
            "Nil."
        )

    else:

        remarks = (
            "The temperature correction of 0.027 % should be applied to the calibration data for each degree difference in the working temperature from the calibration temperature where no temperature compensation is already done (correction will be positive for rise and negative for fall in temperature)."
        )

    print("\n========== RECEIVED FROM FRONTEND ==========")
    print(input_data)
    print("===========================================\n")



    results_layout = dbc.Alert(

        [

            html.H3("ISO 376 Analysis"),

            html.Hr(),

            html.H4("Classification Ranges"),

            dash_table.DataTable(
                

                data=[
                    {
                        "Class": f"Class {r['class']}",
                        "Mode": test_type,
                        "From (kN)": f"{float(r['from_force'])} kN",
                        "To (kN)": f"{float(r['to_force'])} kN",
                        "Uncertainty": r["uncertainty"]
                    }
                    for r in results["class_ranges"]
                ],

                columns=[
                    {"name": "Class", "id": "Class"},
                    {"name": "Mode", "id": "Mode"},
                    {"name": "From", "id": "From (kN)"},
                    {"name": "To", "id": "To (kN)"},
                    {"name": "Uncertainty of Measurement", "id": "Uncertainty"},
                ],

                style_cell={"textAlign": "center"},

                style_header={
                    "backgroundColor": "#003F73",
                    "color": "white",
                    "fontWeight": "bold"
                }

            ),

            html.P(
                f"Zero Error (%) : {results['zero_error']:.4f}"
            ),

            html.P(
                f"Creep (%) : {results['creep']:.4f}"
            ),

            html.Hr(),

            dash_table.DataTable(

                data=results["table_df"].to_dict("records"),

                columns=[
                    {"name": c, "id": c}
                    for c in results["table_df"].columns
                ],

                page_size=20,

                style_table={"overflowX": "auto"},

                style_cell={"textAlign": "center"},

                style_header={
                    "backgroundColor": "#003F73",
                    "color": "white",
                    "fontWeight": "bold"
                }

            )

        ],

        color="success"

    )
    
    certificate_status = ""

    if trigger == "generate_certificate_btn":

        input_data["remarks"] = remarks_box

        if certificate_type == "loadcell":

            filename = generate_loadcell_certificate(
                results,
                input_data
            )

        else:

            filename = generate_proving_ring_certificate(
                results,
                input_data
            )

        certificate_status = dbc.Alert(
            f"Certificate generated successfully.\n{filename}\nSaved to: {OUTPUT_DIR}",
            color="success"
        )

        
    return (


        results_layout,

        remarks,

        False,

        certificate_status,

        {"display": "block"}

    )


# =========================================================
# CLEAR RESULT
# Resets only the Compute/output area (classification result,
# remarks, certificate status, generate-button state) so the
# user can start the next certificate cleanly. This does NOT
# touch the input form fields (Environmental Conditions,
# Calibration Readings, Instrument Details, etc.) and does NOT
# reload the page.
# =========================================================

@app.callback(
    Output("classification_output", "children", allow_duplicate=True),
    Output("remarks_box", "value", allow_duplicate=True),
    Output("generate_certificate_btn", "disabled", allow_duplicate=True),
    Output("certificate_status", "children", allow_duplicate=True),
    Output("remarks_section", "style", allow_duplicate=True),

    Input("clear_result_btn", "n_clicks"),

    prevent_initial_call=True
)
def clear_result(n_clicks):

    return (

        None,

        "",

        True,

        "",

        {"display": "none"}

    )


# =========================================================
# MAIN
# =========================================================

def open_browser(port):
    """Open the app in the user's default browser exactly once."""
    webbrowser.open_new(f"http://127.0.0.1:{port}")


def start_server(port):
    app.run(
        debug=False,
        host="127.0.0.1",
        port=port,
        use_reloader=False,
    )


if __name__ == '__main__':
    try:
        chosen_port = find_free_port(5080)
        logging.info("Starting Force Automation on port %s", chosen_port)

        # Open the browser shortly after the server starts, on a
        # background thread, so it doesn't block server startup and
        # is only ever triggered once.
        threading.Timer(1.25, open_browser, args=(chosen_port,)).start()

        start_server(chosen_port)

    except Exception:
        logging.critical("Fatal error during startup:\n%s", traceback.format_exc())
        raise